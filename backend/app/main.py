"""
Minta — FastAPI backend (auto-hébergé)

Pipeline : faster-whisper (transcription) + pyannote.audio (diarisation)
Persistance : fichiers JSON dans JOB_DIR (un fichier par job)
Audio      : conservé dans AUDIO_DIR pour le player synchronisé

Endpoints :
  GET  /api/health
  POST /api/upload                          → upload audio/vidéo + lancement
  GET  /api/transcribe/{id}                 → statut / résultat
  GET  /api/transcribe/{id}/audio           → stream audio pour le player
  GET  /api/transcribe/{id}/export          → export TXT | SRT | JSON | DOCX
  POST /api/transcribe/{id}/analyze         → déclenche l'analyse IA
  PATCH /api/transcribe/{id}/speakers       → renommer les locuteurs
  PATCH /api/transcribe/{id}/segments       → éditer le texte / commentaires des segments
  POST /api/transcribe/{id}/share           → créer un lien de partage
  GET  /api/share/{token}                   → vue publique d'une transcription
  GET  /api/transcripts                     → liste tous les jobs
  DELETE /api/transcribe/{id}              → supprime un job + audio
  GET  /api/settings
  PATCH /api/settings
"""

import json
import logging
import os
import secrets
import subprocess
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Literal, Optional

import aiofiles
from fastapi import BackgroundTasks, Depends, FastAPI, File, Form, HTTPException, Query, Security, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, PlainTextResponse, StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel as PydanticModel
from pythonjsonlogger import jsonlogger

# ---------------------------------------------------------------------------
# Logging structuré JSON
# ---------------------------------------------------------------------------

LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO").upper()
handler = logging.StreamHandler()
handler.setFormatter(jsonlogger.JsonFormatter("%(asctime)s %(levelname)s %(name)s %(message)s"))
logging.basicConfig(level=LOG_LEVEL, handlers=[handler])
log = logging.getLogger("minta")

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

HF_TOKEN: str = os.getenv("HF_TOKEN", "")
WHISPER_MODEL: str = os.getenv("WHISPER_MODEL", "large-v3")
BASE_DIR = Path(os.getenv("DATA_DIR", "/tmp/minta"))
JOB_DIR = BASE_DIR / "jobs"
UPLOAD_DIR = BASE_DIR / "uploads"
AUDIO_DIR = BASE_DIR / "audio"
SETTINGS_FILE = BASE_DIR / "settings.json"

AI_DEFAULT_PROVIDER: str = os.getenv("AI_DEFAULT_PROVIDER", "anthropic")
ANTHROPIC_API_KEY: str = os.getenv("ANTHROPIC_API_KEY", "")
OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY", "")

# ---------------------------------------------------------------------------
# Auth JWT (optionnel — activé uniquement si AUTH_ENABLED=true)
# ---------------------------------------------------------------------------

AUTH_ENABLED: bool = os.getenv("AUTH_ENABLED", "false").lower() == "true"
AUTH_USERNAME: str = os.getenv("AUTH_USERNAME", "admin")
AUTH_PASSWORD: str = os.getenv("AUTH_PASSWORD", "")  # En clair, hashé au démarrage
JWT_SECRET: str = os.getenv("JWT_SECRET", secrets.token_hex(32))  # généré si absent
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 24

_bearer_scheme = HTTPBearer(auto_error=False)

def _hash_password(pwd: str) -> str:
    from passlib.context import CryptContext
    return CryptContext(schemes=["bcrypt"], deprecated="auto").hash(pwd)

def _verify_password(plain: str, hashed: str) -> bool:
    from passlib.context import CryptContext
    return CryptContext(schemes=["bcrypt"], deprecated="auto").verify(plain, hashed)

def _create_token(username: str) -> str:
    from jose import jwt
    from datetime import timedelta
    expire = datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)
    return jwt.encode({"sub": username, "exp": expire}, JWT_SECRET, algorithm=JWT_ALGORITHM)

def _decode_token(token: str) -> str:
    from jose import JWTError, jwt
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        username: str = payload.get("sub", "")
        if not username:
            raise ValueError("sub manquant")
        return username
    except JWTError as exc:
        raise HTTPException(status_code=401, detail=f"Token invalide : {exc}")

def require_auth(credentials: Optional[HTTPAuthorizationCredentials] = Security(_bearer_scheme)) -> str:
    """Dépendance FastAPI : vérifie le JWT si AUTH_ENABLED, sinon laisse passer."""
    if not AUTH_ENABLED:
        return "anonymous"
    if not credentials:
        raise HTTPException(status_code=401, detail="Token d'authentification requis", headers={"WWW-Authenticate": "Bearer"})
    return _decode_token(credentials.credentials)

# Hash du mot de passe au démarrage (uniquement si auth activé)
_auth_password_hash: str = ""
if AUTH_ENABLED:
    if not AUTH_PASSWORD:
        log.warning("AUTH_ENABLED=true mais AUTH_PASSWORD non défini — authentification désactivée")
        AUTH_ENABLED = False
    else:
        _auth_password_hash = _hash_password(AUTH_PASSWORD)
        log.info("auth enabled", extra={"username": AUTH_USERNAME})

for _d in (JOB_DIR, UPLOAD_DIR, AUDIO_DIR):
    _d.mkdir(parents=True, exist_ok=True)

# Extensions vidéo — ffmpeg extraira la piste audio
VIDEO_EXTENSIONS = {".mp4", ".mkv", ".avi", ".mov", ".m4v", ".wmv", ".flv"}
AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".webm", ".ogg", ".flac", ".aac", ".opus"}

INITIAL_PROMPT = "Transcription d'une réunion professionnelle."

# ---------------------------------------------------------------------------
# Détection device
# ---------------------------------------------------------------------------

def _detect_device() -> tuple[str, str]:
    try:
        import torch
        if torch.cuda.is_available():
            return "cuda", "float16"
    except ImportError:
        pass
    return "cpu", "int8"

DEVICE, COMPUTE_TYPE = _detect_device()
log.info("device detected", extra={"device": DEVICE, "compute_type": COMPUTE_TYPE})

# ---------------------------------------------------------------------------
# Chargement lazy des modèles
# ---------------------------------------------------------------------------

_whisper_model = None
_diarize_pipeline = None
_models_loaded = False
_models_error: Optional[str] = None


def _load_models():
    global _whisper_model, _diarize_pipeline, _models_loaded, _models_error
    if _models_loaded or _models_error:
        return

    try:
        from faster_whisper import WhisperModel
        _whisper_model = WhisperModel(WHISPER_MODEL, device=DEVICE, compute_type=COMPUTE_TYPE)
        log.info("whisper loaded", extra={"model": WHISPER_MODEL})
    except Exception as exc:
        _models_error = f"Chargement Whisper échoué : {exc}"
        log.error(_models_error)
        return

    if not HF_TOKEN:
        log.warning("HF_TOKEN absent — diarisation indisponible")
    else:
        try:
            from pyannote.audio import Pipeline
            _diarize_pipeline = Pipeline.from_pretrained(
                "pyannote/speaker-diarization-3.1",
                use_auth_token=HF_TOKEN,
            )
            if DEVICE == "cuda":
                import torch
                _diarize_pipeline.to(torch.device("cuda"))
            log.info("pyannote loaded")
        except Exception as exc:
            log.warning("pyannote non disponible", extra={"error": str(exc)})

    _models_loaded = True


# ---------------------------------------------------------------------------
# Persistance jobs
# ---------------------------------------------------------------------------

def _job_path(job_id: str) -> Path:
    return JOB_DIR / f"{job_id}.json"


def _save_job(job: dict):
    with open(_job_path(job["id"]), "w", encoding="utf-8") as f:
        json.dump(job, f, ensure_ascii=False, indent=2)


def _load_job(job_id: str) -> dict:
    p = _job_path(job_id)
    if not p.exists():
        raise HTTPException(status_code=404, detail="Job introuvable")
    with open(p, encoding="utf-8") as f:
        return json.load(f)


def _list_jobs() -> list[dict]:
    jobs = []
    for p in sorted(JOB_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            with open(p, encoding="utf-8") as f:
                jobs.append(json.load(f))
        except Exception:
            pass
    return jobs


def _running_jobs_count() -> int:
    return sum(1 for j in _list_jobs() if j.get("status") == "processing")


def _find_audio(job_id: str) -> Optional[Path]:
    for ext in list(AUDIO_EXTENSIONS) + [".wav"]:
        p = AUDIO_DIR / f"{job_id}{ext}"
        if p.exists():
            return p
    return None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _ms(seconds: float) -> int:
    return int(seconds * 1000)


def _srt_ts(ms: int) -> str:
    h = ms // 3_600_000
    m = (ms % 3_600_000) // 60_000
    s = (ms % 60_000) // 1000
    cs = ms % 1000
    return f"{h:02d}:{m:02d}:{s:02d},{cs:03d}"


def _extract_audio(src: Path, dst: Path) -> Path:
    """Extrait la piste audio d'un fichier vidéo en WAV via ffmpeg."""
    result = subprocess.run(
        ["ffmpeg", "-y", "-i", str(src), "-vn", "-ac", "1", "-ar", "16000", "-f", "wav", str(dst)],
        capture_output=True, timeout=300,
    )
    if result.returncode != 0:
        raise RuntimeError(f"ffmpeg extraction échouée : {result.stderr.decode()[:500]}")
    return dst


def _fire_webhook(webhook_url: str, job: dict):
    """Envoie une notification HTTP POST vers le webhook configuré."""
    try:
        import urllib.request
        payload = json.dumps({
            "job_id": job["id"],
            "status": job["status"],
            "filename": job.get("filename"),
            "duration_ms": job.get("duration_ms", 0),
            "speakers": job.get("speakers", []),
            "word_count": job.get("word_count", 0),
        }).encode()
        req = urllib.request.Request(webhook_url, data=payload,
                                      headers={"Content-Type": "application/json"}, method="POST")
        urllib.request.urlopen(req, timeout=10)
        log.info("webhook fired", extra={"url": webhook_url, "job_id": job["id"]})
    except Exception as exc:
        log.warning("webhook failed", extra={"url": webhook_url, "error": str(exc)})


# ---------------------------------------------------------------------------
# Auto-analyse IA post-transcription
# ---------------------------------------------------------------------------

def _trigger_auto_analyze(job: dict, settings: dict):
    """Lance l'analyse IA dans un thread séparé après la transcription."""
    import asyncio, threading
    from .ai.analyzer import run_analysis

    provider = settings.get("default_provider", "anthropic")
    api_key = ANTHROPIC_API_KEY if provider == "anthropic" else OPENAI_API_KEY
    if not api_key:
        log.warning("auto_analyze: clé API absente, analyse ignorée", extra={"provider": provider})
        return
    template = settings.get("auto_analyze_template", "meeting")

    def _run():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(run_analysis(
                job=job,
                save_job_fn=_save_job,
                provider_name=provider,
                api_key=api_key,
                mcp_servers_config=settings.get("mcp_servers", {}),
                active_server_names=[],
                template=template,
            ))
        finally:
            loop.close()

    threading.Thread(target=_run, daemon=True).start()
    log.info("auto_analyze triggered", extra={"job_id": job["id"], "provider": provider})


# ---------------------------------------------------------------------------
# Pipeline transcription
# ---------------------------------------------------------------------------

def _run_pipeline(job_id: str, raw_path: Path, language: str):
    job = _load_job(job_id)
    audio_path = raw_path

    try:
        _load_models()

        if not _whisper_model:
            job["status"] = "error"
            job["error"] = _models_error or "Modèle non chargé"
            _save_job(job)
            return

        log.info("pipeline start", extra={"job_id": job_id})

        # Extraction audio si vidéo
        if raw_path.suffix.lower() in VIDEO_EXTENSIONS:
            wav_path = UPLOAD_DIR / f"{job_id}.wav"
            audio_path = _extract_audio(raw_path, wav_path)
            raw_path.unlink(missing_ok=True)
            log.info("video extracted", extra={"job_id": job_id})

        # Copie vers AUDIO_DIR pour le player
        audio_dest = AUDIO_DIR / f"{job_id}{audio_path.suffix}"
        import shutil
        shutil.copy2(audio_path, audio_dest)

        # Whisper — language=None → auto-détection
        whisper_lang = None if language == "auto" else language
        prompt = INITIAL_PROMPT if not whisper_lang or whisper_lang == "fr" else ""

        segments_iter, info = _whisper_model.transcribe(
            str(audio_path),
            language=whisper_lang,
            initial_prompt=prompt or None,
            word_timestamps=True,
        )
        segments = list(segments_iter)
        detected_language = info.language
        log.info("whisper done", extra={"job_id": job_id, "segments": len(segments), "lang": detected_language})

        # Pyannote (optionnel)
        utterances = []
        if _diarize_pipeline:
            diarization = _diarize_pipeline(str(audio_path))
            speaker_turns = [
                (turn.start, turn.end, speaker)
                for turn, _, speaker in diarization.itertracks(yield_label=True)
            ]
            for seg in segments:
                best_speaker, best_overlap = "SPEAKER_00", 0.0
                for t_start, t_end, spk in speaker_turns:
                    overlap = max(0.0, min(seg.end, t_end) - max(seg.start, t_start))
                    if overlap > best_overlap:
                        best_overlap, best_speaker = overlap, spk
                utterances.append({
                    "speaker": best_speaker,
                    "start": _ms(seg.start), "end": _ms(seg.end),
                    "text": seg.text.strip(),
                    "words": [{"text": w.word, "start": _ms(w.start), "end": _ms(w.end)} for w in (seg.words or [])],
                })
        else:
            for seg in segments:
                utterances.append({
                    "speaker": "SPEAKER_00",
                    "start": _ms(seg.start), "end": _ms(seg.end),
                    "text": seg.text.strip(),
                    "words": [{"text": w.word, "start": _ms(w.start), "end": _ms(w.end)} for w in (seg.words or [])],
                })

        speakers = list(dict.fromkeys(u["speaker"] for u in utterances))
        job.update({
            "status": "completed",
            "completed_at": datetime.utcnow().isoformat(),
            "language": detected_language,
            "utterances": utterances,
            "text": " ".join(u["text"] for u in utterances),
            "speakers": speakers,
            "speaker_names": {s: f"Locuteur {chr(65 + i)}" for i, s in enumerate(speakers)},
            "duration_ms": utterances[-1]["end"] if utterances else 0,
            "word_count": sum(len(u["words"]) for u in utterances),
            "has_audio": True,
        })
        log.info("pipeline done", extra={"job_id": job_id, "speakers": len(speakers)})

    except Exception as exc:
        log.exception("pipeline error", extra={"job_id": job_id})
        job["status"] = "error"
        job["error"] = str(exc)

    finally:
        _save_job(job)
        try:
            audio_path.unlink(missing_ok=True)
        except Exception:
            pass

        # Webhook + Auto-analyse
        settings = _load_settings()
        if wh := settings.get("webhook_url"):
            _fire_webhook(wh, job)
        if job.get("status") == "completed" and settings.get("auto_analyze"):
            _trigger_auto_analyze(job, settings)


# ---------------------------------------------------------------------------
# Export DOCX
# ---------------------------------------------------------------------------

def _export_docx(job: dict, utterances: list, speaker_names: dict, analysis: Optional[dict]):
    import io
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx non installé")

    COLORS = [
        RGBColor(0x0C, 0x44, 0x7C),
        RGBColor(0x3B, 0x6D, 0x11),
        RGBColor(0x85, 0x4F, 0x0B),
        RGBColor(0x72, 0x24, 0x3E),
        RGBColor(0x3C, 0x34, 0x89),
    ]
    speakers = job.get("speakers", [])
    spk_idx = {s: i for i, s in enumerate(speakers)}

    doc = Document()
    doc.core_properties.title = job.get("filename", "Transcription Minta")

    # Titre
    title = doc.add_heading("Transcription — Minta", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Métadonnées
    meta = doc.add_paragraph()
    meta.add_run(f"Fichier : {job.get('filename', '—')}   |   ")
    meta.add_run(f"Langue : {job.get('language', '—')}   |   ")
    meta.add_run(f"Date : {(job.get('completed_at') or '')[:10]}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Résumé IA
    if analysis and analysis.get("status") == "completed":
        if analysis.get("summary"):
            doc.add_heading("Résumé", 1)
            doc.add_paragraph(analysis["summary"])
        if analysis.get("topics"):
            doc.add_heading("Sujets", 2)
            doc.add_paragraph("  •  ".join(analysis["topics"]))
        if analysis.get("decisions"):
            doc.add_heading("Décisions", 2)
            for d in analysis["decisions"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(d)
        if analysis.get("actions"):
            doc.add_heading("Actions", 2)
            for a in analysis["actions"]:
                text = a["text"]
                if a.get("assignee"): text += f" (@{a['assignee']})"
                if a.get("due"): text += f" [échéance: {a['due']}]"
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(text)
        if analysis.get("chapters"):
            doc.add_heading("Chapitres", 2)
            for ch in analysis["chapters"]:
                p = doc.add_paragraph()
                run = p.add_run(f"[{_srt_ts(ch.get('start_ms', 0))[:5]}]  {ch['title']}")
                run.bold = True
                if ch.get("summary"):
                    doc.add_paragraph(ch["summary"])
        doc.add_paragraph()

    # Transcription
    doc.add_heading("Transcription", 1)
    prev_speaker = None
    for u in utterances:
        spk = u["speaker"]
        name = speaker_names.get(spk, spk)
        color = COLORS[spk_idx.get(spk, 0) % len(COLORS)]
        ts = _srt_ts(u["start"])[:8]

        if spk != prev_speaker:
            p = doc.add_paragraph()
            run = p.add_run(f"{name}  —  {ts}")
            run.bold = True
            run.font.color.rgb = color
            run.font.size = Pt(10)
            prev_speaker = spk

        p = doc.add_paragraph(u["text"])
        p.paragraph_format.left_indent = Pt(18)
        if u.get("comment"):
            note = doc.add_paragraph(f"  💬 {u['comment']}")
            note.paragraph_format.left_indent = Pt(24)
            run = note.runs[0]
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = job.get("filename", job["id"])
    filename = Path(filename).stem + ".docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# App FastAPI
# ---------------------------------------------------------------------------

app = FastAPI(title="Minta API", version="0.5.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", os.getenv("FRONTEND_ORIGIN", "")],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "app": "Minta",
        "version": "0.6.0",
        "device": DEVICE,
        "compute_type": COMPUTE_TYPE,
        "whisper_model": WHISPER_MODEL,
        "models_loaded": _models_loaded,
        "models_error": _models_error,
        "hf_token_configured": bool(HF_TOKEN),
        "diarization_available": _diarize_pipeline is not None,
        "jobs_running": _running_jobs_count(),
        "auth_enabled": AUTH_ENABLED,
    }


# ---------------------------------------------------------------------------
# Auth endpoints
# ---------------------------------------------------------------------------

class LoginRequest(PydanticModel):
    username: str
    password: str


@app.post("/api/auth/login")
def login(req: LoginRequest):
    """Retourne un JWT si les identifiants sont corrects. Toujours disponible."""
    if not AUTH_ENABLED:
        return {"token": None, "auth_enabled": False}
    if req.username != AUTH_USERNAME or not _verify_password(req.password, _auth_password_hash):
        raise HTTPException(status_code=401, detail="Identifiants invalides")
    token = _create_token(req.username)
    log.info("login success", extra={"username": req.username})
    return {"token": token, "auth_enabled": True, "expires_in": JWT_EXPIRE_HOURS * 3600}


@app.post("/api/upload", status_code=202)
async def upload_audio(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("auto"),
    _: str = Depends(require_auth),
):
    """Reçoit un fichier audio ou vidéo, crée un job et lance le pipeline."""
    job_id = str(uuid.uuid4())
    suffix = Path(file.filename or "audio").suffix.lower() or ".webm"
    raw_path = UPLOAD_DIR / f"{job_id}{suffix}"

    async with aiofiles.open(raw_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    running = _running_jobs_count()
    job = {
        "id": job_id,
        "filename": file.filename,
        "language": language,
        "status": "processing",
        "created_at": datetime.utcnow().isoformat(),
        "completed_at": None,
        "utterances": [],
        "text": "",
        "speakers": [],
        "speaker_names": {},
        "duration_ms": 0,
        "word_count": 0,
        "error": None,
        "queue_position": running,
        "has_audio": False,
        "share_token": None,
    }
    _save_job(job)

    background_tasks.add_task(_run_pipeline, job_id, raw_path, language)
    log.info("job created", extra={"job_id": job_id, "filename": file.filename, "language": language})
    return {
        "id": job_id,
        "status": "processing",
        "queue_position": running,
        "message": (
            "Traitement en cours, votre fichier sera traité après les jobs en cours."
            if running > 0 else "Traitement démarré."
        ),
    }


@app.get("/api/transcribe/{job_id}")
def get_job(job_id: str, _: str = Depends(require_auth)):
    return _load_job(job_id)


@app.get("/api/transcribe/{job_id}/audio")
def get_audio(job_id: str, _: str = Depends(require_auth)):
    _load_job(job_id)  # vérifie que le job existe
    audio = _find_audio(job_id)
    if not audio:
        raise HTTPException(status_code=404, detail="Fichier audio non disponible")
    media_type = "audio/wav" if audio.suffix == ".wav" else f"audio/{audio.suffix.lstrip('.')}"
    return FileResponse(audio, media_type=media_type, filename=audio.name)


@app.get("/api/transcribe/{job_id}/export")
def export_job(job_id: str, format: Literal["txt", "srt", "json", "docx"] = Query("txt"), _: str = Depends(require_auth)):
    job = _load_job(job_id)
    if job["status"] != "completed":
        raise HTTPException(status_code=409, detail="Transcription non terminée")

    utterances = job.get("utterances", [])
    speaker_names: dict = job.get("speaker_names", {})
    analysis = job.get("analysis")

    def spk_label(spk: str) -> str:
        return speaker_names.get(spk, spk)

    if format == "json":
        return job

    if format == "txt":
        lines = []
        if analysis and analysis.get("status") == "completed":
            if analysis.get("summary"):
                lines += ["=== RÉSUMÉ ===", analysis["summary"], ""]
            if analysis.get("decisions"):
                lines += ["=== DÉCISIONS ==="] + [f"• {d}" for d in analysis["decisions"]] + [""]
            if analysis.get("actions"):
                lines += ["=== ACTIONS ==="] + [
                    f"• {a['text']}" + (f" (@{a['assignee']})" if a.get("assignee") else "") +
                    (f" [échéance: {a['due']}]" if a.get("due") else "")
                    for a in analysis["actions"]
                ] + [""]
            lines.append("=== TRANSCRIPTION ===")
        for u in utterances:
            ts = _srt_ts(u["start"]).replace(",", ".").rsplit(".", 1)[0]
            comment = f"  [{u['comment']}]" if u.get("comment") else ""
            lines.append(f"[{ts}] [{spk_label(u['speaker'])}] {u['text']}{comment}")
        content = "\n".join(lines)
        return PlainTextResponse(content, media_type="text/plain; charset=utf-8",
                                 headers={"Content-Disposition": f'attachment; filename="{job_id}.txt"'})

    if format == "srt":
        blocks = []
        for i, u in enumerate(utterances, 1):
            blocks.append(f"{i}\n{_srt_ts(u['start'])} --> {_srt_ts(u['end'])}\n[{spk_label(u['speaker'])}] {u['text']}\n")
        return PlainTextResponse("\n".join(blocks), media_type="text/plain; charset=utf-8",
                                 headers={"Content-Disposition": f'attachment; filename="{job_id}.srt"'})

    if format == "docx":
        return _export_docx(job, utterances, speaker_names, analysis)


class SpeakerNamesUpdate(PydanticModel):
    speaker_names: dict[str, str]


@app.patch("/api/transcribe/{job_id}/speakers")
def update_speaker_names(job_id: str, body: SpeakerNamesUpdate, _: str = Depends(require_auth)):
    job = _load_job(job_id)
    job["speaker_names"] = {**job.get("speaker_names", {}), **body.speaker_names}
    _save_job(job)
    return {"speaker_names": job["speaker_names"]}


class SegmentPatch(PydanticModel):
    index: int
    text: Optional[str] = None
    comment: Optional[str] = None


class SegmentsPatch(PydanticModel):
    segments: list[SegmentPatch]


@app.patch("/api/transcribe/{job_id}/segments")
def patch_segments(job_id: str, body: SegmentsPatch, _: str = Depends(require_auth)):
    """Édite le texte et/ou le commentaire d'un ou plusieurs segments."""
    job = _load_job(job_id)
    utterances = job.get("utterances", [])
    for patch in body.segments:
        if 0 <= patch.index < len(utterances):
            if patch.text is not None:
                utterances[patch.index]["text"] = patch.text
                utterances[patch.index]["edited"] = True
            if patch.comment is not None:
                utterances[patch.index]["comment"] = patch.comment
    job["utterances"] = utterances
    # Recalculer le texte complet
    job["text"] = " ".join(u["text"] for u in utterances)
    _save_job(job)
    return {"utterances": utterances}


@app.post("/api/transcribe/{job_id}/share")
def create_share(job_id: str, _: str = Depends(require_auth)):
    job = _load_job(job_id)
    if job["status"] != "completed":
        raise HTTPException(status_code=409, detail="Transcription non terminée")
    if not job.get("share_token"):
        job["share_token"] = secrets.token_urlsafe(24)
        _save_job(job)
    return {"share_token": job["share_token"], "url": f"/share/{job['share_token']}"}


@app.get("/api/share/{token}")
def get_shared(token: str):
    for job in _list_jobs():
        if job.get("share_token") == token:
            return {k: job[k] for k in ("id", "filename", "language", "utterances", "text",
                                         "speakers", "speaker_names", "duration_ms", "word_count",
                                         "completed_at") if k in job}
    raise HTTPException(status_code=404, detail="Lien de partage introuvable ou expiré")


@app.get("/api/transcripts")
def list_transcripts(
    q: Optional[str] = Query(None, description="Recherche full-text dans les transcriptions"),
    tag: Optional[str] = Query(None, description="Filtrer par tag"),
    _: str = Depends(require_auth),
):
    jobs = _list_jobs()

    if q:
        needle = q.lower()
        def _matches(j: dict) -> bool:
            if needle in (j.get("filename") or "").lower():
                return True
            if needle in j.get("text", "").lower():
                return True
            names = " ".join(j.get("speaker_names", {}).values()).lower()
            if needle in names:
                return True
            analysis = j.get("analysis") or {}
            if needle in analysis.get("summary", "").lower():
                return True
            if any(needle in d.lower() for d in analysis.get("decisions", [])):
                return True
            if any(needle in t.lower() for t in analysis.get("topics", [])):
                return True
            return False
        jobs = [j for j in jobs if _matches(j)]

    if tag:
        jobs = [j for j in jobs if tag in j.get("tags", [])]

    return [
        {
            "id": j["id"],
            "filename": j.get("filename"),
            "language": j.get("language"),
            "status": j["status"],
            "created_at": j.get("created_at"),
            "completed_at": j.get("completed_at"),
            "duration_ms": j.get("duration_ms", 0),
            "speakers": j.get("speakers", []),
            "speaker_names": j.get("speaker_names", {}),
            "word_count": j.get("word_count", 0),
            "error": j.get("error"),
            "has_audio": j.get("has_audio", False),
            "share_token": j.get("share_token"),
            "tags": j.get("tags", []),
            "analysis_status": j.get("analysis", {}).get("status") if j.get("analysis") else None,
        }
        for j in jobs
    ]


class TagsUpdate(PydanticModel):
    tags: list[str]


@app.patch("/api/transcribe/{job_id}/tags")
def update_tags(job_id: str, body: TagsUpdate, _: str = Depends(require_auth)):
    job = _load_job(job_id)
    job["tags"] = [t.strip() for t in body.tags if t.strip()]
    _save_job(job)
    return {"tags": job["tags"]}


@app.delete("/api/transcribe/{job_id}", status_code=204)
def delete_job(job_id: str, _: str = Depends(require_auth)):
    p = _job_path(job_id)
    if not p.exists():
        raise HTTPException(status_code=404, detail="Job introuvable")
    p.unlink()
    audio = _find_audio(job_id)
    if audio:
        audio.unlink(missing_ok=True)
    log.info("job deleted", extra={"job_id": job_id})


# ---------------------------------------------------------------------------
# Settings
# ---------------------------------------------------------------------------

_DEFAULT_SETTINGS: dict = {
    "mcp_servers": {},
    "default_provider": AI_DEFAULT_PROVIDER,
    "webhook_url": "",
    "auto_analyze": False,
    "auto_analyze_template": "meeting",
}


def _load_settings() -> dict:
    if SETTINGS_FILE.exists():
        with open(SETTINGS_FILE, encoding="utf-8") as f:
            saved = json.load(f)
        return {**_DEFAULT_SETTINGS, **saved}
    return dict(_DEFAULT_SETTINGS)


def _save_settings(settings: dict):
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)


@app.get("/api/settings")
def get_settings(_: str = Depends(require_auth)):
    settings = _load_settings()
    settings["providers_available"] = {
        "anthropic": bool(ANTHROPIC_API_KEY),
        "openai": bool(OPENAI_API_KEY),
    }
    return settings


class SettingsPatch(PydanticModel):
    mcp_servers: Optional[dict] = None
    default_provider: Optional[str] = None
    webhook_url: Optional[str] = None
    auto_analyze: Optional[bool] = None
    auto_analyze_template: Optional[str] = None


@app.patch("/api/settings")
def patch_settings(patch: SettingsPatch, _: str = Depends(require_auth)):
    settings = _load_settings()
    if patch.mcp_servers is not None:
        settings["mcp_servers"] = patch.mcp_servers
    if patch.default_provider is not None:
        if patch.default_provider not in ("anthropic", "openai"):
            raise HTTPException(status_code=400, detail="Fournisseur inconnu")
        settings["default_provider"] = patch.default_provider
    if patch.webhook_url is not None:
        settings["webhook_url"] = patch.webhook_url
    if patch.auto_analyze is not None:
        settings["auto_analyze"] = patch.auto_analyze
    if patch.auto_analyze_template is not None:
        settings["auto_analyze_template"] = patch.auto_analyze_template
    _save_settings(settings)
    return settings


# ---------------------------------------------------------------------------
# Analyse IA
# ---------------------------------------------------------------------------

class AnalyzeRequest(PydanticModel):
    provider: str = "anthropic"
    api_key: Optional[str] = None
    mcp_servers: list[str] = []
    template: str = "meeting"


async def _run_analysis_task(job_id: str, provider: str, api_key: str, active_servers: list[str], template: str):
    from .ai.analyzer import run_analysis
    job = _load_job(job_id)
    settings = _load_settings()
    await run_analysis(
        job=job,
        save_job_fn=_save_job,
        provider_name=provider,
        api_key=api_key,
        mcp_servers_config=settings.get("mcp_servers", {}),
        active_server_names=active_servers,
        template=template,
    )


@app.post("/api/transcribe/{job_id}/analyze", status_code=202)
async def analyze_job(job_id: str, req: AnalyzeRequest, background_tasks: BackgroundTasks, _: str = Depends(require_auth)):
    job = _load_job(job_id)
    if job["status"] != "completed":
        raise HTTPException(status_code=409, detail="Transcription non terminée")

    provider = req.provider
    api_key = req.api_key or (ANTHROPIC_API_KEY if provider == "anthropic" else OPENAI_API_KEY)
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail=f"Clé API {provider} absente. Fournissez api_key dans la requête ou configurez la variable d'environnement.",
        )

    background_tasks.add_task(_run_analysis_task, job_id, provider, api_key, req.mcp_servers, req.template)
    log.info("analysis requested", extra={"job_id": job_id, "provider": provider, "template": req.template})
    return {"job_id": job_id, "analysis_status": "running"}


# ---------------------------------------------------------------------------
# WebSocket — transcription temps réel
# ---------------------------------------------------------------------------

@app.websocket("/api/ws/transcribe")
async def ws_transcribe(
    websocket: WebSocket,
    language: str = "auto",
    token: Optional[str] = None,
):
    """
    Transcription en temps réel pendant l'enregistrement.

    Protocole client → serveur :
      binary frame  : blob audio WebM complet (tous les chunks depuis le début)
      text frame    : {"type": "stop"}

    Protocole serveur → client :
      {"type": "partial", "utterances": [...], "language": "fr"}
      {"type": "error",   "message": "..."}
      {"type": "stopped"}
    """
    # Auth JWT via query param (les WS ne supportent pas les headers custom)
    if AUTH_ENABLED:
        if not token:
            await websocket.close(code=1008, reason="Token requis")
            return
        try:
            _decode_token(token)
        except HTTPException:
            await websocket.close(code=1008, reason="Token invalide")
            return

    await websocket.accept()
    log.info("ws_transcribe connected")

    import asyncio

    async def _transcribe_blob(audio_bytes: bytes) -> list[dict]:
        """Transcrit un blob audio en utilisant un executor pour ne pas bloquer la boucle."""
        tmp_webm = Path(tempfile.mktemp(suffix=".webm", dir=str(UPLOAD_DIR)))
        tmp_wav  = tmp_webm.with_suffix(".wav")
        try:
            tmp_webm.write_bytes(audio_bytes)
            # Conversion WebM → WAV 16kHz mono
            try:
                _extract_audio(tmp_webm, tmp_wav)
                audio_path = tmp_wav
            except Exception:
                audio_path = tmp_webm

            lang = None if language == "auto" else language

            def _run_whisper():
                segs, info = _whisper_model.transcribe(  # type: ignore[union-attr]
                    str(audio_path),
                    language=lang,
                    word_timestamps=False,
                    beam_size=1,
                    vad_filter=True,
                    condition_on_previous_text=False,
                )
                return list(segs), info

            loop = asyncio.get_event_loop()
            segments, info = await loop.run_in_executor(None, _run_whisper)

            utterances = [
                {
                    "speaker": "SPEAKER_00",
                    "start": _ms(seg.start),
                    "end": _ms(seg.end),
                    "text": seg.text.strip(),
                    "words": [],
                }
                for seg in segments if seg.text.strip()
            ]
            return utterances, getattr(info, "language", language or "?")
        finally:
            tmp_webm.unlink(missing_ok=True)
            tmp_wav.unlink(missing_ok=True)

    try:
        _load_models()
        if not _whisper_model:
            await websocket.send_json({"type": "error", "message": "Modèle Whisper non chargé"})
            await websocket.close()
            return

        while True:
            message = await websocket.receive()

            if message["type"] == "websocket.disconnect":
                break

            # Blob audio binaire
            if message.get("bytes"):
                try:
                    utterances, detected_lang = await _transcribe_blob(message["bytes"])
                    await websocket.send_json({
                        "type": "partial",
                        "utterances": utterances,
                        "language": detected_lang,
                    })
                except Exception as exc:
                    log.warning("ws_transcribe chunk error", extra={"error": str(exc)})
                    await websocket.send_json({"type": "error", "message": str(exc)})

            # Message texte (commande)
            elif message.get("text"):
                try:
                    data = json.loads(message["text"])
                except json.JSONDecodeError:
                    continue
                if data.get("type") == "stop":
                    await websocket.send_json({"type": "stopped"})
                    break

    except WebSocketDisconnect:
        pass
    except Exception as exc:
        log.exception("ws_transcribe error")
        try:
            await websocket.send_json({"type": "error", "message": str(exc)})
        except Exception:
            pass
    finally:
        log.info("ws_transcribe disconnected")
